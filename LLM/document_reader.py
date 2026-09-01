"""
LLM/document_reader.py

Extension-first triage for /api/upload-document.

The old endpoint decided how to read a file purely from the browser-
supplied Content-Type, and only accepted image/jpeg, image/png and
application/pdf -- anything else (a .docx, .csv, .txt, .xlsx...) was
rejected outright with a 400, even though none of those need OCR at
all since their text is already digital.

This module fixes that by looking at the file's *extension* first and
routing accordingly:

  - image / pdf extensions  -> OCR path. Callers should hand the bytes
    to LLM.extract_document_text(), which already does native-PDF-text
    first and falls back to GPT-4o Vision only for scanned pages.
  - any other supported extension (.txt, .md, .csv, .tsv, .json,
    .docx, .xlsx, .xlsm) -> direct_read path, handled entirely here.
    No vision/OCR call is made -- there's nothing to "see", the text
    is just parsed straight out of the file.
  - anything else -> unsupported; caller should reject with a clear
    list of what *is* supported instead of guessing.

Both paths return the same shape:
    {"text": str, "page_count": int, "pages_read": int, "method": str}
so the rest of the app (document_store, the doc_context injected into
chat, etc.) never needs to know which path was taken.
"""
import csv
import io
import json
import logging
import os

logger = logging.getLogger("document-reader")

IMAGE_EXTENSIONS = {".png", ".jpg", ".jpeg", ".webp", ".gif", ".bmp", ".tiff", ".tif"}
PDF_EXTENSIONS = {".pdf"}
OCR_EXTENSIONS = IMAGE_EXTENSIONS | PDF_EXTENSIONS

TEXT_EXTENSIONS = {".txt", ".md", ".rtf"}
CSV_EXTENSIONS = {".csv", ".tsv"}
JSON_EXTENSIONS = {".json"}
DOCX_EXTENSIONS = {".docx"}
XLSX_EXTENSIONS = {".xlsx", ".xlsm"}

DIRECT_READ_EXTENSIONS = (
    TEXT_EXTENSIONS | CSV_EXTENSIONS | JSON_EXTENSIONS | DOCX_EXTENSIONS | XLSX_EXTENSIONS
)

ALL_SUPPORTED_EXTENSIONS = OCR_EXTENSIONS | DIRECT_READ_EXTENSIONS


def get_extension(filename: str) -> str:
    return os.path.splitext(filename or "")[1].lower()


def classify_extension(filename: str) -> str:
    """Returns 'ocr', 'direct', or 'unsupported' based on the file's
    extension alone -- this is the FIRST thing the upload endpoint
    should check, before touching content_type or file bytes."""
    ext = get_extension(filename)
    if ext in OCR_EXTENSIONS:
        return "ocr"
    if ext in DIRECT_READ_EXTENSIONS:
        return "direct"
    return "unsupported"


def extract_text_generic(file_bytes: bytes, filename: str) -> dict:
    """Reads a NON-image/PDF file straight off its digital text -- no
    OCR/vision call needed since nothing here was scanned. Mirrors the
    return shape of LLM.extract_document_text() so callers can treat
    both extraction paths identically:
        {"text": str, "page_count": int, "pages_read": int, "method": "direct_read"}

    Raises RuntimeError on failure (bad zip for .docx/.xlsx, undecodable
    bytes, etc.) -- callers should catch this and turn it into an HTTP
    error, same as OCR failures already are.
    """
    ext = get_extension(filename)

    try:
        if ext in TEXT_EXTENSIONS:
            text = file_bytes.decode("utf-8", errors="replace")
            return {"text": text, "page_count": 1, "pages_read": 1, "method": "direct_read"}

        if ext in CSV_EXTENSIONS:
            delimiter = "\t" if ext == ".tsv" else ","
            decoded = file_bytes.decode("utf-8", errors="replace")
            rows = list(csv.reader(io.StringIO(decoded), delimiter=delimiter))
            if not rows:
                return {"text": "", "page_count": 1, "pages_read": 1, "method": "direct_read"}
            header, body = rows[0], rows[1:]
            lines = [
                "| " + " | ".join(header) + " |",
                "| " + " | ".join(["---"] * len(header)) + " |",
            ]
            lines += ["| " + " | ".join(row) + " |" for row in body]
            return {"text": "\n".join(lines), "page_count": 1, "pages_read": 1, "method": "direct_read"}

        if ext in JSON_EXTENSIONS:
            decoded = file_bytes.decode("utf-8", errors="replace")
            try:
                text = json.dumps(json.loads(decoded), indent=2, ensure_ascii=False)
            except Exception:
                text = decoded  # not valid JSON -- fall back to raw text rather than failing
            return {"text": text, "page_count": 1, "pages_read": 1, "method": "direct_read"}

        if ext in DOCX_EXTENSIONS:
            import docx  # python-docx

            document = docx.Document(io.BytesIO(file_bytes))
            parts = [p.text for p in document.paragraphs if p.text.strip()]
            for table in document.tables:
                for row in table.rows:
                    parts.append(" | ".join(cell.text.strip() for cell in row.cells))
            return {"text": "\n".join(parts), "page_count": 1, "pages_read": 1, "method": "direct_read"}

        if ext in XLSX_EXTENSIONS:
            import openpyxl

            wb = openpyxl.load_workbook(io.BytesIO(file_bytes), data_only=True, read_only=True)
            sheet_texts = []
            for sheet in wb.worksheets:
                lines = [f"--- Sheet: {sheet.title} ---"]
                for row in sheet.iter_rows(values_only=True):
                    cells = ["" if c is None else str(c) for c in row]
                    if any(cells):
                        lines.append(" | ".join(cells))
                sheet_texts.append("\n".join(lines))
            return {
                "text": "\n\n".join(sheet_texts),
                "page_count": len(wb.worksheets),
                "pages_read": len(wb.worksheets),
                "method": "direct_read",
            }

        raise RuntimeError(f"Unsupported extension for direct read: '{ext}'")

    except RuntimeError:
        raise
    except Exception as e:
        logger.exception("Error extracting text from '%s'", filename)
        raise RuntimeError(f"Could not read '{filename}': {e}")
